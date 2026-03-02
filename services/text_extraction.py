"""
Reusable text extraction service.
Works with file bytes (from R2) rather than file paths on disk.
Reuses patterns from quotation_processor.py.
"""
import io
import fitz  # PyMuPDF
import anthropic
import base64
import xml.etree.ElementTree as ET
from config import claude_api_key

_anthropic_client = None


def _get_anthropic_client() -> anthropic.Anthropic:
    global _anthropic_client
    if _anthropic_client is None:
        _anthropic_client = anthropic.Anthropic(api_key=claude_api_key)
    return _anthropic_client


def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using PyMuPDF, with Vision fallback for image-based pages."""
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    pages_text = []

    for page_num, page in enumerate(doc):
        text = page.get_text().strip()
        if text and len(text) > 30:
            pages_text.append(text)
        else:
            # Page has little/no text — likely scanned/image-based, use Vision
            try:
                pix = page.get_pixmap(dpi=200)
                img_bytes = pix.tobytes("png")
                ocr_text = extract_text_from_image_bytes(img_bytes, "image/png")
                if ocr_text and ocr_text.strip():
                    pages_text.append(ocr_text)
            except Exception as e:
                print(f"[TextExtraction] Vision fallback failed for page {page_num}: {e}")

    doc.close()
    return "\n".join(pages_text)


def extract_text_from_image_bytes(file_bytes: bytes, media_type: str) -> str:
    """Extract text from image bytes using Claude Vision API."""
    client = _get_anthropic_client()
    image_base64 = base64.b64encode(file_bytes).decode('utf-8')

    message = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=20000,
        temperature=0,
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "image",
                    "source": {"type": "base64", "media_type": media_type, "data": image_base64}
                },
                {
                    "type": "text",
                    "text": "Please extract ALL text from this image. Return only the extracted text, maintaining structure."
                }
            ]
        }]
    )
    return message.content[0].text.strip()


def extract_text_from_txt_bytes(file_bytes: bytes) -> str:
    """Extract text from .txt file bytes."""
    try:
        return file_bytes.decode('utf-8').strip()
    except UnicodeDecodeError:
        return file_bytes.decode('latin-1').strip()


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    """Extract text from .docx file bytes using python-docx."""
    from docx import Document

    doc = Document(io.BytesIO(file_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)

    return "\n".join(paragraphs)


def extract_text_from_xml_cfdi(file_bytes: bytes) -> str:
    """Extract structured text from a CFDI XML invoice."""
    try:
        text_content = file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        text_content = file_bytes.decode('latin-1')

    try:
        root = ET.fromstring(text_content)
    except ET.ParseError:
        # If XML parsing fails, return raw text
        return text_content

    ns = {
        'cfdi': 'http://www.sat.gob.mx/cfd/4',
        'tfd': 'http://www.sat.gob.mx/TimbreFiscalDigital',
    }

    lines = ["FACTURA CFDI"]

    # Comprobante attributes
    fecha = root.attrib.get('Fecha', '')
    subtotal = root.attrib.get('SubTotal', '')
    total = root.attrib.get('Total', '')
    moneda = root.attrib.get('Moneda', '')
    forma_pago = root.attrib.get('FormaPago', '')
    metodo_pago = root.attrib.get('MetodoPago', '')

    if fecha:
        lines.append(f"Fecha: {fecha}")
    if total:
        lines.append(f"Total: ${total} {moneda}")
    if subtotal:
        lines.append(f"Subtotal: ${subtotal} {moneda}")
    if forma_pago:
        lines.append(f"Forma de Pago: {forma_pago}")
    if metodo_pago:
        lines.append(f"Método de Pago: {metodo_pago}")

    # Emisor
    emisor = root.find('cfdi:Emisor', ns)
    if emisor is not None:
        rfc_emisor = emisor.attrib.get('Rfc', '')
        nombre_emisor = emisor.attrib.get('Nombre', '')
        regimen = emisor.attrib.get('RegimenFiscal', '')
        lines.append(f"\nEmisor: {nombre_emisor}")
        lines.append(f"RFC Emisor: {rfc_emisor}")
        if regimen:
            lines.append(f"Régimen Fiscal: {regimen}")

    # Receptor
    receptor = root.find('cfdi:Receptor', ns)
    if receptor is not None:
        rfc_receptor = receptor.attrib.get('Rfc', '')
        nombre_receptor = receptor.attrib.get('Nombre', '')
        uso_cfdi = receptor.attrib.get('UsoCFDI', '')
        lines.append(f"\nReceptor: {nombre_receptor}")
        lines.append(f"RFC Receptor: {rfc_receptor}")
        if uso_cfdi:
            lines.append(f"Uso CFDI: {uso_cfdi}")

    # Conceptos
    conceptos = root.find('cfdi:Conceptos', ns)
    if conceptos is not None:
        lines.append("\nConceptos:")
        for concepto in conceptos.findall('cfdi:Concepto', ns):
            desc = concepto.attrib.get('Descripcion', '')
            cantidad = concepto.attrib.get('Cantidad', '')
            unidad = concepto.attrib.get('Unidad', '')
            valor = concepto.attrib.get('ValorUnitario', '')
            importe = concepto.attrib.get('Importe', '')
            lines.append(f"  - {desc}: {cantidad} {unidad} x ${valor} = ${importe}")

    # Impuestos
    impuestos = root.find('cfdi:Impuestos', ns)
    if impuestos is not None:
        total_traslados = impuestos.attrib.get('TotalImpuestosTrasladados', '')
        if total_traslados:
            lines.append(f"\nIVA/Impuestos: ${total_traslados}")

    # Timbre Fiscal
    complemento = root.find('cfdi:Complemento', ns)
    if complemento is not None:
        timbre = complemento.find('tfd:TimbreFiscalDigital', ns)
        if timbre is not None:
            uuid = timbre.attrib.get('UUID', '')
            fecha_timbrado = timbre.attrib.get('FechaTimbrado', '')
            if uuid:
                lines.append(f"\nUUID: {uuid}")
            if fecha_timbrado:
                lines.append(f"Fecha Timbrado: {fecha_timbrado}")

    return "\n".join(lines)


MEDIA_TYPE_MAP = {
    'image/png': 'image/png',
    'image/jpeg': 'image/jpeg',
    'image/gif': 'image/gif',
    'image/bmp': 'image/bmp',
    'image/tiff': 'image/tiff',
    'image/webp': 'image/webp',
}

DOCX_CONTENT_TYPES = {
    'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    'application/msword',
}

XML_CONTENT_TYPES = {
    'application/xml',
    'text/xml',
}


def extract_text(file_bytes: bytes, content_type: str, original_filename: str) -> str:
    """Main extraction dispatcher. Returns extracted text or raises an exception."""
    lower_filename = original_filename.lower()

    # DOCX files
    if content_type in DOCX_CONTENT_TYPES or lower_filename.endswith('.docx'):
        return extract_text_from_docx_bytes(file_bytes)

    # XML files (CFDI invoices)
    if content_type in XML_CONTENT_TYPES or lower_filename.endswith('.xml'):
        return extract_text_from_xml_cfdi(file_bytes)

    # PDF files
    if content_type == 'application/pdf':
        return extract_text_from_pdf_bytes(file_bytes)

    # Image files
    if content_type in MEDIA_TYPE_MAP:
        return extract_text_from_image_bytes(file_bytes, MEDIA_TYPE_MAP[content_type])

    # Plain text files
    if content_type.startswith('text/') or lower_filename.endswith('.txt'):
        return extract_text_from_txt_bytes(file_bytes)

    raise ValueError(f"Unsupported content type for text extraction: {content_type}")


def is_extractable(content_type: str, original_filename: str) -> bool:
    """Check whether we can extract text from this file type."""
    lower_filename = original_filename.lower()

    if content_type == 'application/pdf':
        return True
    if content_type in MEDIA_TYPE_MAP:
        return True
    if content_type.startswith('text/') or lower_filename.endswith('.txt'):
        return True
    if content_type in DOCX_CONTENT_TYPES or lower_filename.endswith('.docx'):
        return True
    if content_type in XML_CONTENT_TYPES or lower_filename.endswith('.xml'):
        return True
    return False
